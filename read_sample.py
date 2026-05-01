from docx import Document

doc = Document(r"C:\Users\Admin\OneDrive\Desktop\DAMA_Report_v2_2026-04-13T06-32-26.docx")

for para in doc.paragraphs:
    print(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in table.rows[0].cells:
            print(cell.text, end=" | ")
        print()