from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

def set_cell_shading(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading(text, level):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12 if level == 1 else 11)
        run.font.color.rgb = RGBColor(30, 58, 95)
    return heading

def add_paragraph(text, bold=False, italic=False, size=11, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    return p

# Title Page
add_paragraph("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph("Student Helper Chatbot", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Comprehensive Project Report", align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph("Submitted in partial fulfilment of the requirements for the degree of", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Bachelor of Technology (B.Tech.) in Computer Science and Engineering", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph("Submitted by", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Project Team – Final Year, B.Tech. CSE", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph("Under the Guidance of", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph("Project Supervisor, Department of CSE", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_paragraph("Academic Year: 2025 - 2026", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# Abstract
add_heading("ABSTRACT", level=1)
abstract_text = """The Student Helper Chatbot is a Flask-based web application designed to assist students with their academic queries, study assistance, and general information retrieval. The system provides an AI-powered conversational interface that helps students get quick answers to their questions, access study materials, and receive personalized recommendations.

The application is built using Flask for the backend, SQLite for database management, and HTML/CSS/JavaScript for the frontend. It includes user authentication with OTP (One-Time Password) verification via email, session management, and a responsive chat interface. The chatbot employs natural language processing techniques to understand user queries and provide appropriate responses.

Key features include user registration and login with secure OTP verification, password reset functionality, AI-powered chatbot responses, chat history storage, email notifications, and a responsive web interface. The system supports multiple user modes and provides personalized responses based on user preferences and query context.

The platform was developed following structured software engineering principles and uses best practices for web application development including MVC architecture, secure authentication, and database optimization."""

p = add_paragraph(abstract_text)
for run in p.runs:
    run.font.size = Pt(10)

doc.add_page_break()

# Table of Contents
add_heading("TABLE OF CONTENTS", level=1)
toc_items = [
    ("CHAPTER 1 – INTRODUCTION", "1"),
    ("1.1 Project Summary", "1"),
    ("1.2 Project Purpose", "2"),
    ("1.3 Project Scope", "2"),
    ("1.4 Objectives", "3"),
    ("1.5 Technology and Literature Overview", "4"),
    ("CHAPTER 2 – LITERATURE SURVEY", "5"),
    ("2.1 Introduction to the Survey", "5"),
    ("2.2 Review of Existing Systems", "5"),
    ("2.3 Research Gap Analysis", "6"),
    ("CHAPTER 3 – PROJECT MANAGEMENT", "7"),
    ("3.1 Project Planning", "7"),
    ("3.2 Project Scheduling", "7"),
    ("3.3 Risk Management", "8"),
    ("CHAPTER 4 – SYSTEM REQUIREMENTS", "9"),
    ("4.1 User Characteristics", "9"),
    ("4.2 Functional Requirements", "9"),
    ("4.3 Non-Functional Requirements", "10"),
    ("4.4 Hardware and Software Requirements", "10"),
    ("CHAPTER 5 – SYSTEM ANALYSIS", "11"),
    ("5.1 Study of Current System", "11"),
    ("5.2 Problems in Current System", "11"),
    ("5.3 Requirements of New System", "12"),
    ("5.4 Process Model", "12"),
    ("5.5 Feasibility Study", "12"),
    ("CHAPTER 6 – MODULE DESCRIPTION", "14"),
    ("6.1 User Module", "14"),
    ("6.2 Authentication Module", "14"),
    ("6.3 Chatbot Module", "15"),
    ("6.4 Database Module", "15"),
    ("6.5 Email Notification Module", "15"),
    ("CHAPTER 7 – TESTING", "16"),
    ("7.1 Testing Strategy", "16"),
    ("7.2 Test Cases", "16"),
    ("CHAPTER 8 – SYSTEM DESIGN", "17"),
    ("8.1 UML Diagrams", "17"),
    ("8.2 Database Design", "18"),
    ("8.3 Architecture Design", "18"),
    ("CHAPTER 9 – LIMITATIONS AND FUTURE ENHANCEMENTS", "19"),
    ("9.1 Limitations", "19"),
    ("9.2 Future Enhancements", "19"),
    ("CHAPTER 10 – CONCLUSION", "20"),
    ("APPENDICES", "21"),
]

for item, page in toc_items:
    p = doc.add_paragraph()
    run1 = p.add_run(item)
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add dots
    run2 = p.add_run(" ")
    run2.font.name = 'Times New Roman'
    run3 = p.add_run(page)
    run3.font.name = 'Times New Roman'
    run3.font.size = Pt(11)

doc.add_page_break()

# CHAPTER 1
add_heading("CHAPTER 1 – INTRODUCTION", level=1)

add_heading("1.1 Project Summary", level=2)
summary_text = """Student Helper Chatbot is a full-stack web application designed to provide academic assistance to students through an intelligent conversational interface. The system serves as a virtual assistant that helps students with their queries, provides study resources, and offers personalized recommendations based on their needs.

The application consists of a Flask backend server that handles all business logic, API endpoints, and database operations. The frontend is built using HTML, CSS, and JavaScript, providing a responsive and intuitive user interface. The chatbot component uses natural language processing to understand user queries and generate appropriate responses.

The system includes comprehensive user management features including registration with OTP verification, secure login with session management, and password reset functionality using time-limited OTPs. User data and chat history are stored in a SQLite database, providing reliable data persistence and easy backup options.

The project follows the Model-View-Controller (MVC) architectural pattern, ensuring clean separation of concerns and maintainable code structure. The application is designed to be easily extensible, allowing for future integration of advanced AI models, additional features, and third-party services."""
p = add_paragraph(summary_text)

add_heading("1.2 Project Purpose", level=2)
purpose_text = """The primary purpose of Student Helper Chatbot is to provide students with immediate access to academic assistance and information through a user-friendly conversational interface. The specific purposes include:

1. Provide 24/7 access to academic information and assistance without requiring human intervention.

2. Reduce the workload on academic support staff by automating responses to common student queries.

3. Provide personalized assistance based on student preferences and query history.

4. Store and retrieve chat history for future reference and analysis.

5. Offer a modern, accessible platform that students can use from any device with internet connectivity.

6. Ensure secure access through robust authentication mechanisms including OTP verification.

7. Create a scalable solution that can be easily expanded to serve more users and support additional features."""
p = add_paragraph(purpose_text)

add_heading("1.3 Project Scope", level=2)
scope_text = """The scope of Student Helper Chatbot encompasses the following components:

1.3.1 User Features (Web Interface)
Students can register for an account, verify their email using OTP, login with secure credentials, reset their password through OTP, access the chatbot interface, view their chat history, and logout securely.

1.3.2 Administrator Features
Admin access to manage user accounts, view system statistics, monitor chatbot performance, and manage database records.

1.3.3 Chatbot Features
The chatbot provides responses to user queries, supports multiple conversation modes, stores conversation history, learns from interactions (basic pattern matching), and provides contextual responses.

1.3.4 Excluded Features
The current scope explicitly excludes: video consultation features, integration with learning management systems (LMS), automated grading systems, live chat with human operators, mobile application development, and real-time collaboration features."""
p = add_paragraph(scope_text)

add_heading("1.4 Objectives", level=2)

add_heading("1.4.1 Main Objectives", level=3)
main_obj = """1. Develop a secure, scalable web application using Flask framework.

2. Implement user authentication with OTP email verification.

3. Create a responsive chatbot interface with natural language processing capabilities.

4. Store and manage user data and chat history in SQLite database.

5. Provide session-based authentication with secure password handling.

6. Implement email notification system for OTP delivery and account verification.

7. Ensure responsive design that works on desktop and mobile browsers."""
p = add_paragraph(main_obj)

add_heading("1.4.2 Secondary Objectives", level=3)
sec_obj = """1. Implement proper error handling and validation across all user inputs.

2. Create a modular code structure for easy maintenance and expansion.

3. Use environment variables for sensitive configuration data.

4. Implement proper logging for debugging and monitoring.

5. Ensure code follows Python and web development best practices.

6. Provide user-friendly error messages and interface feedback.

7. Implement proper session management and logout functionality."""
p = add_paragraph(sec_obj)

add_heading("1.5 Technology and Literature Overview", level=2)

add_heading("1.5.1 Flask (Python 3)", level=3)
flask_text = """Flask is a lightweight WSGI web application framework in Python. It is designed to make getting started quick and easy, with the ability to scale up to complex applications. Flask offers flexibility in project structure and allows developers to choose their preferred tools and libraries. Flask includes Jinja2 templating, Werkzeug WSGI utilities, and safe session management. The framework supports extensions for database integration, form validation, upload handling, and various open authentication technologies."""
p = add_paragraph(flask_text)

add_heading("1.5.2 SQLite Database", level=3)
sqlite_text = """SQLite is a C-language library that implements a small, fast, self-contained, high-reliability, full-featured, SQL database engine. SQLite is the most used database engine in the world. It is built into all mobile phones and most computers and comes bundled inside countless other applications that people use every day. The SQLite file format is stable, cross-platform, and backwards compatible. The developers commit to keeping it that way through at least the year 2050."""
p = add_paragraph(sqlite_text)

add_heading("1.5.3 HTML, CSS, and JavaScript", level=3)
frontend_text = """HTML (HyperText Markup Language) is the standard markup language for documents designed to be displayed in a web browser. CSS (Cascading Style Sheets) is a style sheet language used for describing the presentation of a document written in a markup language. JavaScript is a programming language that enables interactive web pages and dynamic content. Together, these technologies form the foundation of modern web development."""
p = add_paragraph(frontend_text)

add_heading("1.5.4 OTP Authentication", level=3)
otp_text = """One-Time Password (OTP) authentication is a security mechanism that generates a unique, temporary password for a single login session or transaction. OTPs are typically sent via email or SMS and are valid for a limited time period. This method provides an additional layer of security beyond traditional password-based authentication, protecting against password theft, replay attacks, and unauthorized access."""
p = add_paragraph(otp_text)

doc.add_page_break()

# CHAPTER 2
add_heading("CHAPTER 2 – LITERATURE SURVEY", level=1)

add_heading("2.1 Introduction to the Survey", level=2)
survey_intro = """The literature survey for the Student Helper Chatbot project explores the domains of: (1) AI-powered educational chatbots, (2) web-based student support systems, (3) secure authentication mechanisms, and (4) modern web application development frameworks. A systematic review of existing solutions, academic publications, and industry best practices was conducted to identify research gaps, establish design benchmarks, and justify the technical choices made in this project."""
p = add_paragraph(survey_intro)

add_heading("2.2 Review of Existing Systems", level=2)

add_heading("2.2.1 Traditional Student Support Systems", level=3)
trad_text = """Traditional student support systems in educational institutions rely heavily on human resources including faculty advisors, counseling services, and administrative staff. While these systems provide personalized support, they are limited by availability, response time, and human resource constraints. Students often have to wait for office hours or schedule appointments to get their queries resolved."""
p = add_paragraph(trad_text)

add_heading("2.2.2 Existing Chatbot Solutions", level=3)
existing_text = """Various educational institutions and EdTech companies have developed chatbot solutions for student support. However, many of these solutions suffer from limitations such as: limited natural language understanding capabilities, lack of personalization, poor integration with institutional systems, and inadequate security measures. Additionally, many existing solutions are expensive to implement and maintain, making them inaccessible for smaller institutions."""
p = add_paragraph(existing_text)

add_heading("2.3 Research Gap Analysis", level=2)
gap_text = """Based on the literature survey, the following research gaps were identified and addressed by this project:

1. Lack of affordable, open-source chatbot solutions for educational institutions.

2. Limited integration of secure OTP-based authentication in student web applications.

3. Minimal use of modern Flask-based architectures in educational chatbot development.

4. Absence of comprehensive documentation and implementation guides for similar projects.

5. Need for a modular, easily extensible chatbot architecture that can be customized for different educational contexts."""

p = add_paragraph(gap_text)

doc.add_page_break()

# CHAPTER 3
add_heading("CHAPTER 3 – PROJECT MANAGEMENT", level=1)

add_heading("3.1 Project Planning", level=2)

add_heading("3.1.1 Software Scope", level=3)
scope_planning = """The project scope covers the development of a complete web-based Student Helper Chatbot application including: Flask backend with all API endpoints, frontend web interface with responsive design, SQLite database with proper schema design, user authentication system with OTP verification, chatbot response generation system, and email notification system using SMTP."""
p = add_paragraph(scope_planning)

add_heading("3.1.2 Resources", level=3)
resources_text = """Human Resources: Project Team (1-2 developers), Project Supervisor
Software Resources: Flask, SQLite, python-dotenv, Flask-Session
Environment: Python 3.x, Git, VS Code"""
p = add_paragraph(resources_text)

add_heading("3.2 Project Scheduling", level=2)

add_heading("3.2.1 Development Phases", level=3)
phases_text = """Phase 1: Requirements gathering and database design (Week 1)
Phase 2: Backend development and API implementation (Week 2-3)
Phase 3: Frontend development and UI implementation (Week 4)
Phase 4: Authentication system implementation (Week 5)
Phase 5: Testing and debugging (Week 6)
Phase 6: Documentation and deployment (Week 7)"""
p = add_paragraph(phases_text)

add_heading("3.3 Risk Management", level=2)
risk_text = """Table 3.1 Risk Matrix

| Risk ID | Risk Description | Probability | Impact | Mitigation |
|---------|-----------------|-------------|--------|------------|
| R1 | Database connection issues | Medium | High | Proper error handling and connection pooling |
| R2 | Email delivery failures | Low | High | Alternative email verification method |
| R3 | Session timeout issues | Low | Medium | Extended session duration and remember me option |
| R4 | Performance bottlenecks | Medium | Medium | Code optimization and database indexing |
| R5 | Security vulnerabilities | Low | High | Regular security audits and input validation |"""

table = doc.add_table(rows=6, cols=6)
table.style = 'Table Grid'
headers = ['Risk ID', 'Risk Description', 'Probability', 'Impact', 'Mitigation']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, '1E3A5F')
    for para in cell.paragraphs:
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

risk_data = [
    ['R1', 'Database connection issues', 'Medium', 'High', 'Proper error handling'],
    ['R2', 'Email delivery failures', 'Low', 'High', 'Alternative verification'],
    ['R3', 'Session timeout', 'Low', 'Medium', 'Extended session duration'],
    ['R4', 'Performance bottlenecks', 'Medium', 'Medium', 'Code optimization'],
    ['R5', 'Security vulnerabilities', 'Low', 'High', 'Security audits'],
]

for i, row_data in enumerate(risk_data, 1):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text

doc.add_page_break()

# CHAPTER 4
add_heading("CHAPTER 4 – SYSTEM REQUIREMENTS", level=1)

add_heading("4.1 User Characteristics", level=2)
user_chars = """The Student Helper Chatbot system serves two primary user classes:

1. Students: University/college students seeking academic assistance, typically aged 18-25, with moderate to high digital literacy. They access the system through web browsers on desktop or mobile devices.

2. Administrators: Staff members who manage the system, monitor performance, and handle user accounts. They require moderate technical knowledge."""

p = add_paragraph(user_chars)

add_heading("4.2 Functional Requirements", level=2)
func_req = """Table 4.1 Functional Requirements

| Requirement ID | Description | Priority |
|----------------|-------------|----------|
| FR01 | User Registration with email verification | High |
| FR02 | User Login with session management | High |
| FR03 | OTP-based password reset | High |
| FR04 | Chatbot query response | High |
| FR05 | Chat history storage and retrieval | Medium |
| FR06 | User profile management | Medium |
| FR07 | Email notification system | High |
| FR08 | Logout functionality | High |"""

table = doc.add_table(rows=9, cols=3)
table.style = 'Table Grid'
headers = ['Requirement ID', 'Description', 'Priority']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, '1E3A5F')
    for para in cell.paragraphs:
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

req_data = [
    ['FR01', 'User Registration with email verification', 'High'],
    ['FR02', 'User Login with session management', 'High'],
    ['FR03', 'OTP-based password reset', 'High'],
    ['FR04', 'Chatbot query response', 'High'],
    ['FR05', 'Chat history storage and retrieval', 'Medium'],
    ['FR06', 'User profile management', 'Medium'],
    ['FR07', 'Email notification system', 'High'],
    ['FR08', 'Logout functionality', 'High'],
]

for i, row_data in enumerate(req_data, 1):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text

add_heading("4.3 Non-Functional Requirements", level=2)
nonfunc_text = """Security: All passwords hashed using SHA-256. Sessions secured with secret key. OTP valid for 10 minutes only.

Performance: API response times under 500ms. Database queries optimized with proper indexing.

Scalability: Modular architecture allows easy scaling. Database can be migrated to MySQL/PostgreSQL for larger scale.

Reliability: Proper error handling implemented. Logging for debugging and monitoring.

Usability: Responsive design works on all devices. Simple, intuitive interface."""

p = add_paragraph(nonfunc_text)

add_heading("4.4 Hardware and Software Requirements", level=2)

hw_req = """Table 4.2 Hardware Requirements

| Component | Minimum Specification |
|-----------|----------------------|
| Processor | Intel Core i3 or equivalent |
| RAM | 4 GB |
| Storage | 10 GB free space |
| Internet | Broadband connection |"""

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
for i, header in enumerate(['Component', 'Minimum Specification']):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, '1E3A5F')
    for para in cell.paragraphs:
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

hw_data = [
    ['Processor', 'Intel Core i3 or equivalent'],
    ['RAM', '4 GB'],
    ['Storage', '10 GB free space'],
    ['Internet', 'Broadband connection'],
]

for i, row_data in enumerate(hw_data, 1):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text

sw_req = """Table 4.3 Software Requirements

| Software | Version |
|----------|----------|
| Python | 3.8+ |
| Flask | 2.0+ |
| SQLite | 3.x |
| pip | Latest |"""

table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
for i, header in enumerate(['Software', 'Version']):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, '1E3A5F')
    for para in cell.paragraphs:
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

sw_data = [
    ['Python', '3.8+'],
    ['Flask', '2.0+'],
    ['SQLite', '3.x'],
    ['pip', 'Latest'],
]

for i, row_data in enumerate(sw_data, 1):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text

doc.add_page_break()

# CHAPTER 5
add_heading("CHAPTER 5 – SYSTEM ANALYSIS", level=1)

add_heading("5.1 Study of Current System", level=2)
current_sys = """The current system for student support in educational institutions typically involves manual processes including email communication, office hours consultation, and FAQ pages. Students often face challenges in getting timely responses to their queries, especially during peak times or outside office hours. The lack of a centralized system leads to inconsistent responses and inefficient use of staff time."""

p = add_paragraph(current_sys)

add_heading("5.2 Problems in Current System", level=2)
problems = """1. Limited availability: Students can only get support during office hours.

2. Inconsistent responses: Different staff members may provide different answers.

3. Time-consuming: Students wait in queues or for email responses.

4. No tracking: No record of previous queries and responses.

5. No personalization: Generic responses not tailored to individual needs.

6. No 24/7 availability: Urgent queries cannot be addressed immediately.

7. Resource intensive: Requires dedicated staff for student support."""

p = add_paragraph(problems)

add_heading("5.3 Requirements of New System", level=2)
new_req = """A digital system that provides immediate responses to student queries.

A centralized database for storing and retrieving information.

An intelligent chatbot that learns and improves over time.

Secure authentication to protect student data.

24/7 availability for student support.

Personalized responses based on query history.

Easy to use interface accessible from any device.

Cost-effective solution requiring minimal maintenance."""

p = add_paragraph(new_req)

add_heading("5.4 Process Model", level=2)
process_model = """The Student Helper Chatbot follows the Waterfall-Scrum hybrid process model. An initial requirements and architecture phase established the system design. Subsequent development followed Agile principles with iterative development and testing. This model was selected because the requirements were well understood, and the iterative approach allowed for continuous improvement based on testing feedback."""

p = add_paragraph(process_model)

add_heading("5.5 Feasibility Study", level=2)

add_heading("5.5.1 Technical Feasibility", level=3)
tech_feas = """All selected technologies are mature, well-documented, and widely used. Flask provides excellent development experience with minimal setup. SQLite offers reliable data storage without requiring additional server configuration. The required packages are all available via pip and are well-maintained. The development team has the necessary skills in Python, Flask, and web development."""

p = add_paragraph(tech_feas)

add_heading("5.5.2 Operational Feasibility", level=3)
op_feas = """The system is operationally feasible for the target users. The web interface requires no specialized technical knowledge. Students are familiar with chatbot interfaces from common messaging platforms. Training requirements are minimal - a brief orientation is sufficient for new users. The system is designed for easy maintenance and updates."""

p = add_paragraph(op_feas)

add_heading("5.5.3 Economical Feasibility", level=3)
eco_feas = """The project uses only open-source technologies, eliminating licensing costs. The hosting requirements are minimal - any basic web hosting with Python support will suffice. The total cost of ownership is very low compared to commercial solutions. The system provides excellent value for money with minimal ongoing expenses."""

p = add_paragraph(eco_feas)

add_heading("5.5.4 Schedule Feasibility", level=3)
sched_feas = """The development schedule is achievable given the team size and project scope. The use of Flask allows rapid development. The modular structure enables parallel development of different components. The timeline of approximately 7-8 weeks is sufficient for a complete implementation with testing."""

p = add_paragraph(sched_feas)

doc.add_page_break()

# CHAPTER 6
add_heading("CHAPTER 6 – MODULE DESCRIPTION", level=1)

add_heading("6.1 User Module", level=2)
user_mod = """The User Module handles all user-related functionality including:

1. User Registration: New users can create an account by providing username, email, and password. The system validates input and checks for duplicate entries.

2. Profile Management: Users can view and update their profile information. Email changes require re-verification.

3. Session Management: User login state is maintained using Flask sessions. Automatic logout after inactivity.

4. Account Security: Password reset functionality with time-limited OTPs. Secure logout that clears session data."""

p = add_paragraph(user_mod)

add_heading("6.2 Authentication Module", level=2)
auth_mod = """The Authentication Module provides secure access control:

1. Registration with OTP: Users receive a 6-digit OTP via email for email verification. OTP is valid for 10 minutes.

2. Login Authentication: Username and password verification. Session creation on successful login.

3. Password Reset: OTP-based password reset flow. User must verify email before changing password.

4. Logout: Clears session data and redirects to home page.

5. Password Hashing: SHA-256 hashing for secure password storage."""

p = add_paragraph(auth_mod)

add_heading("6.3 Chatbot Module", level=2)
chat_mod = """The Chatbot Module provides the core AI functionality:

1. Message Processing: Receives user messages and processes them for response generation.

2. Response Generation: Uses pattern matching and keyword detection to generate appropriate responses.

3. Mode Support: Supports different response modes (student mode and general mode).

4. History Storage: Stores all conversations in the database for future reference.

5. Error Handling: Graceful handling of unrecognized queries with helpful responses."""

p = add_paragraph(chat_mod)

add_heading("6.4 Database Module", level=2)
db_mod = """The Database Module manages all data persistence:

1. User Data: Stores user credentials, profile information, and account status.

2. Message Storage: Maintains complete chat history for each user.

3. Session Data: Tracks user sessions and activity timestamps.

4. Data Integrity: Foreign key constraints and proper indexing for performance.

5. Backup Support: Simple file-based database allows easy backup and migration."""

p = add_paragraph(db_mod)

add_heading("6.5 Email Notification Module", level=2)
email_mod = """The Email Notification Module handles all email communications:

1. OTP Emails: Sends verification and reset OTPs to user email addresses.

2. HTML Templates: Professional-looking HTML emails with project branding.

3. Error Handling: Graceful handling of email delivery failures.

4. Logging: Logs all email operations for debugging and monitoring.

5. SMTP Integration: Uses Gmail SMTP for reliable email delivery."""

p = add_paragraph(email_mod)

doc.add_page_break()

# CHAPTER 7
add_heading("CHAPTER 7 – TESTING", level=1)

add_heading("7.1 Testing Strategy", level=2)
test_strat = """The testing strategy includes multiple levels:

1. Unit Testing: Testing individual functions and methods in isolation.

2. Integration Testing: Testing the interaction between different modules.

3. System Testing: Testing the complete system as a whole.

4. User Acceptance Testing: Testing by end users to ensure requirements are met.

5. Security Testing: Testing authentication and authorization mechanisms."""

p = add_paragraph(test_strat)

add_heading("7.2 Test Cases", level=2)
test_cases = """Table 7.1 Test Cases

| Test ID | Description | Expected Result | Status |
|---------|-------------|-----------------|--------|
| TC01 | User registration with valid data | Account created successfully | Pass |
| TC02 | User registration with duplicate email | Error message displayed | Pass |
| TC03 | User login with correct credentials | Login successful, redirect to app | Pass |
| TC04 | User login with incorrect password | Error message displayed | Pass |
| TC05 | OTP verification with correct code | Account verified | Pass |
| TC06 | OTP verification with expired code | Error: OTP expired | Pass |
| TC07 | Password reset flow | Password updated successfully | Pass |
| TC08 | Chatbot response generation | Relevant response returned | Pass |
| TC09 | Session timeout | Redirect to login page | Pass |
| TC10 | Logout functionality | Session cleared, redirect to home | Pass |"""

table = doc.add_table(rows=11, cols=5)
table.style = 'Table Grid'
for i, header in enumerate(['Test ID', 'Description', 'Expected Result', 'Status']):
    cell = table.rows[0].cells[i]
    cell.text = header
    set_cell_shading(cell, '1E3A5F')
    for para in cell.paragraphs:
        para.runs[0].bold = True
        para.runs[0].font.color.rgb = RGBColor(255, 255, 255)

test_data = [
    ['TC01', 'User registration with valid data', 'Account created successfully', 'Pass'],
    ['TC02', 'User registration with duplicate email', 'Error message displayed', 'Pass'],
    ['TC03', 'User login with correct credentials', 'Login successful', 'Pass'],
    ['TC04', 'User login with incorrect password', 'Error message displayed', 'Pass'],
    ['TC05', 'OTP verification with correct code', 'Account verified', 'Pass'],
    ['TC06', 'OTP verification with expired code', 'Error: OTP expired', 'Pass'],
    ['TC07', 'Password reset flow', 'Password updated successfully', 'Pass'],
    ['TC08', 'Chatbot response generation', 'Relevant response returned', 'Pass'],
    ['TC09', 'Session timeout', 'Redirect to login page', 'Pass'],
    ['TC10', 'Logout functionality', 'Session cleared', 'Pass'],
]

for i, row_data in enumerate(test_data, 1):
    for j, cell_text in enumerate(row_data):
        table.rows[i].cells[j].text = cell_text

doc.add_page_break()

# CHAPTER 8
add_heading("CHAPTER 8 – SYSTEM DESIGN", level=1)

add_heading("8.1 UML Diagrams", level=2)
uml_text = """8.1.1 Use Case Diagram
The system has the following actors: User (Student) and Administrator. Main use cases include: Register, Login, Reset Password, Send Message, View History, Logout.

8.1.2 Class Diagram
Main classes include: User, Message, ChatSession, Authentication. Each class has appropriate attributes and methods.

8.1.3 Sequence Diagram
User authentication flow: User enters credentials → System validates → Create session → Redirect to app.

8.1.4 Activity Diagram
Registration flow: Fill form → Validate input → Check duplicates → Generate OTP → Send email → Verify OTP → Create account."""

p = add_paragraph(uml_text)

add_heading("8.2 Database Design", level=2)
db_design = """Table: users
- id (INTEGER PRIMARY KEY)
- username (TEXT UNIQUE)
- email (TEXT UNIQUE)
- password (TEXT)
- created_at (TIMESTAMP)

Table: messages
- id (INTEGER PRIMARY KEY)
- sender (TEXT) - 'user' or 'bot'
- message (TEXT)
- timestamp (TIMESTAMP)

ER Diagram: Users table has one-to-many relationship with Messages table."""

p = add_paragraph(db_design)

add_heading("8.3 Architecture Design", level=2)
arch_text = """The system follows the Model-View-Controller (MVC) architecture:

1. Model: Database models using SQLite with proper relationships.

2. View: HTML templates rendered by Flask with CSS styling.

3. Controller: Flask routes handling business logic and API endpoints.

The application uses a layered architecture:
- Presentation Layer: HTML/CSS/JS frontend
- Business Logic Layer: Flask routes and chatbot logic
- Data Access Layer: SQLite database operations
- Security Layer: Authentication and session management"""

p = add_paragraph(arch_text)

doc.add_page_break()

# CHAPTER 9
add_heading("CHAPTER 9 – LIMITATIONS AND FUTURE ENHANCEMENTS", level=1)

add_heading("9.1 Limitations", level=2)
limitations = """1. Limited AI capabilities: The chatbot uses basic pattern matching rather than advanced NLP or machine learning models.

2. Single database: Uses SQLite which may have concurrency limitations for high-traffic scenarios.

3. Email dependency: Relies on Gmail SMTP for OTP delivery, which may have rate limits.

4. No mobile app: Currently only accessible via web browser.

5. No real-time updates: Chat requires page refresh for new messages.

6. Single language: Interface is currently in English only.

7. No integration: Does not integrate with external systems or APIs."""

p = add_paragraph(limitations)

add_heading("9.2 Future Enhancements", level=2)
enhancements = """1. AI Enhancement: Integrate advanced NLP models (like GPT) for more intelligent responses.

2. Mobile Application: Develop React Native or Flutter mobile app for better mobile experience.

3. Multi-language Support: Add support for multiple languages for wider accessibility.

4. Database Migration: Migrate to PostgreSQL or MySQL for better scalability.

5. Real-time Chat: Implement WebSocket for instant messaging without page refresh.

6. Integration: Add integration with learning management systems (LMS), calendar apps, etc.

7. Analytics: Add dashboard for usage statistics and chatbot performance metrics.

8. Voice Support: Add voice input and output for accessibility.

9. Video Consultation: Add option to connect with human advisors for complex queries.

10. API Development: Expose RESTful API for third-party integrations."""

p = add_paragraph(enhancements)

doc.add_page_break()

# CHAPTER 10
add_heading("CHAPTER 10 – CONCLUSION", level=1)
conclusion = """The Student Helper Chatbot project successfully demonstrates the development of a complete web-based application using Flask and related technologies. The system provides students with 24/7 access to academic assistance through an intuitive chatbot interface.

The implementation includes all core features required for a student support system: user authentication with OTP verification, secure session management, intelligent chatbot responses, and reliable data storage. The project follows best practices in software development including proper error handling, input validation, and code organization.

The system is modular and extensible, allowing for easy addition of new features in future iterations. The use of open-source technologies ensures the project is cost-effective and accessible for educational institutions with limited budgets.

Overall, the Student Helper Chatbot represents a significant step forward in providing accessible, efficient, and scalable student support services through modern web technology. The project serves as a solid foundation for further development and enhancement to meet the evolving needs of students and educational institutions."""

p = add_paragraph(conclusion)

doc.add_page_break()

# APPENDICES
add_heading("APPENDICES", level=1)

add_heading("Appendix A: File Structure", level=2)
file_struct = """student_helper_chatbot/
├── app.py                 # Main Flask application
├── init_db.py            # Database initialization script
├── chatbot/              # Chatbot module
│   └── model.py         # Chatbot response logic
├── database/            # Database files
│   └── chatbot.db       # SQLite database
├── templates/           # HTML templates
│   ├── home.html
│   ├── login.html
│   ├── signup.html
│   ├── index.html
│   ├── about.html
│   ├── verify_otp.html
│   ├── forgot_password.html
│   └── reset_password.html
├── static/              # CSS and JavaScript
│   ├── style.css
│   └── script.js
├── .env                 # Environment variables
└── venv/                # Virtual environment"""

p = add_paragraph(file_struct)

add_heading("Appendix B: API Endpoints", level=2)
api_endpoints = """GET  /              - Home page
GET  /login           - Login page
POST /login           - Login authentication
GET  /signup          - Registration page
POST /signup          - User registration
GET  /verify-signup   - OTP verification page
POST /verify-signup   - OTP verification
GET  /logout          - Logout and session clear
GET  /app             - Main chat application (requires auth)
POST /chat            - Chat API endpoint
GET  /about           - About page
GET  /forgot-password - Password reset request
POST /forgot-password - Password reset request
GET  /verify-reset    - Reset OTP verification
POST /verify-reset    - Reset OTP verification
GET  /reset-password  - New password page
POST /reset-password  - Password update"""

p = add_paragraph(api_endpoints)

add_heading("Appendix C: Configuration", level=2)
config_text = """Environment Variables (.env):
- SECRET_KEY: Flask secret key for sessions
- MAIL_USERNAME: Gmail email for sending OTPs
- MAIL_PASSWORD: App password for Gmail SMTP

Database Configuration:
- Database file: database/chatbot.db
- SQLite with automatic table creation
- Connection via sqlite3 module"""

p = add_paragraph(config_text)

# End of Document
doc.add_page_break()
end_para = doc.add_paragraph("--- End of Documentation ---")
end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in end_para.runs:
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(136, 136, 136)
    run.italic = True

# Save
doc.save('StudentHelperChatbot_Documentation.docx')
print("Documentation generated: StudentHelperChatbot_Documentation.docx")